"""Layout-aware PDF and document text extraction.

Extraction strategy (PDF):
1. Pull words with geometry + font info via pdfplumber.
2. Detect a true column gutter (sidebar layouts) — only split when a
   persistent vertical gap divides the page; otherwise single flow.
3. Group words into lines top-to-bottom, words left-to-right.
4. Insert a " | " separator for wide horizontal gaps so downstream
   parsers can split right-aligned dates/locations from headings.
5. Clean glyph artifacts ((cid:NNN) icons, ligatures, odd bullets).
6. Collect hyperlink annotations (LinkedIn/GitHub URLs are often link
   targets whose visible text is just "LinkedIn").
"""

from __future__ import annotations

import io
import re
import statistics
from dataclasses import dataclass, field
from pathlib import Path

from jobpilot.services.resume.document import ResumeDocument, ResumeLine

# Wide-gap separator threshold (PDF points). Justified text stretches
# inter-word gaps to ~8-10pt; right-aligned dates/locations leave 30pt+.
_GAP_MIN_PT = 16.0

_CID_RE = re.compile(r"\(cid:\d+\)")

_CHAR_FIXES = str.maketrans(
    {
        "ﬀ": "ff",
        "ﬁ": "fi",
        "ﬂ": "fl",
        "ﬃ": "ffi",
        "ﬄ": "ffl",
        "‘": "'",
        "’": "'",
        "“": '"',
        "”": '"',
        " ": " ",
        "​": "",
        "": "•",
        "●": "•",
        "▪": "•",
        "·": "•",
        "‧": "•",
    }
)


@dataclass
class ExtractedContent:
    text: str
    lines: list[ResumeLine]
    hyperlinks: list[str] = field(default_factory=list)


def extract_from_bytes(content: bytes, filename: str) -> ExtractedContent:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(content)
    if suffix in {".docx", ".doc"}:
        if suffix == ".doc":
            raise ValueError("Legacy .doc files are not supported; please upload PDF or DOCX")
        return _extract_docx(content)
    if suffix in {".txt", ".md"}:
        text = _clean_text(content.decode("utf-8", errors="ignore"))
        doc = ResumeDocument.from_text(text)
        return ExtractedContent(text=doc.raw_text, lines=doc.lines)
    raise ValueError(f"Unsupported file type: {suffix or 'unknown'}. Use PDF, DOCX, or TXT.")


def extract_text_from_bytes(content: bytes, filename: str) -> str:
    return extract_from_bytes(content, filename).text


def _clean_text(text: str) -> str:
    text = text.lstrip("﻿").replace("\r\n", "\n").replace("\r", "\n")
    text = _CID_RE.sub(" ", text)
    text = text.translate(_CHAR_FIXES)
    return text


# ---------------------------------------------------------------- DOCX


def _extract_docx(content: bytes) -> ExtractedContent:
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(io.BytesIO(content))

    lines: list[ResumeLine] = []

    def add_line(text: str, is_header: bool = False) -> None:
        # Tabs align title/date or company/location columns → treat as separators
        text = _clean_text(text).replace("\t", " | ")
        # A paragraph can hold soft line breaks — emit each as its own line
        for piece in text.split("\n"):
            piece = re.sub(r"\s*\|\s*", " | ", piece).strip(" |").strip()
            piece = re.sub(r"[ ]{2,}", " ", piece)
            if piece:
                lines.append(ResumeLine(text=piece, index=len(lines), is_header=is_header))

    # Walk body elements in document order so tables interleave correctly.
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            for para in doc.paragraphs:
                if para._p is child:
                    style = (para.style.name or "").lower() if para.style else ""
                    add_line(para.text, is_header=style.startswith(("heading", "title")))
                    break
        elif child.tag == qn("w:tbl"):
            for table in doc.tables:
                if table._tbl is child:
                    for row in table.rows:
                        cells = [c.text.strip() for c in row.cells if c.text.strip()]
                        # Dedupe merged cells that repeat text
                        seen: list[str] = []
                        for c in cells:
                            if not seen or seen[-1] != c:
                                seen.append(c)
                        if seen:
                            add_line(" | ".join(seen))
                    break

    hyperlinks: list[str] = []
    for rel in doc.part.rels.values():
        if "hyperlink" in rel.reltype and rel.is_external:
            hyperlinks.append(rel.target_ref)

    text = "\n".join(ln.text for ln in lines)
    return ExtractedContent(text=text, lines=lines, hyperlinks=hyperlinks)


# ----------------------------------------------------------------- PDF


def _extract_pdf(content: bytes) -> ExtractedContent:
    import pdfplumber

    all_lines: list[ResumeLine] = []
    hyperlinks: list[str] = []
    sizes: list[float] = []

    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            for link in page.hyperlinks or []:
                uri = link.get("uri")
                if uri and uri not in hyperlinks:
                    hyperlinks.append(uri)

            words = page.extract_words(
                x_tolerance=1.5,
                y_tolerance=3,
                keep_blank_chars=False,
                extra_attrs=["size", "fontname"],
            )
            page_lines = _lines_from_words(words, page.width or 612.0)
            if not page_lines:
                fallback = page.extract_text(x_tolerance=1.5, y_tolerance=3) or ""
                page_lines = [
                    _Line(text=t.strip(), top=0.0, size=None)
                    for t in fallback.splitlines()
                    if t.strip()
                ]
            for pl in page_lines:
                text = _clean_text(pl.text).strip(" |")
                text = re.sub(r"\s*\|\s*", " | ", text)
                text = re.sub(r"[ \t]{2,}", " ", text).strip()
                if not text:
                    continue
                if pl.size:
                    sizes.append(pl.size)
                all_lines.append(
                    ResumeLine(
                        text=text,
                        index=len(all_lines),
                        font_size=pl.size,
                    )
                )

    median_size = statistics.median(sizes) if sizes else None
    for ln in all_lines:
        ln.is_header = _looks_like_header(ln, median_size)

    text = "\n".join(ln.text for ln in all_lines)
    return ExtractedContent(text=text, lines=all_lines, hyperlinks=hyperlinks)


def _looks_like_header(line: ResumeLine, median_size: float | None) -> bool:
    text = line.stripped
    word_count = len(text.split())
    if (
        median_size
        and line.font_size
        and line.font_size >= median_size * 1.15
        and word_count <= 8
    ):
        return True
    return text.isupper() and 3 <= len(text) <= 40 and word_count <= 6


@dataclass
class _Line:
    text: str
    top: float
    size: float | None


def _lines_from_words(words: list[dict], page_width: float) -> list[_Line]:
    if not words:
        return []

    columns = _split_columns(words, page_width)
    result: list[_Line] = []
    for column_words in columns:
        result.extend(_flow_lines(column_words))
    return result


def _split_columns(words: list[dict], page_width: float) -> list[list[dict]]:
    """Split words into columns only when a persistent vertical gutter exists.

    Most resumes are single-column; a false split scrambles reading order,
    so require strong evidence: a gap band in the central page region that
    almost no words cross, with substantial content on both sides.
    """
    if len(words) < 20:
        return [words]

    bin_w = 4.0
    n_bins = int(page_width / bin_w) + 1
    coverage = [0] * n_bins
    for w in words:
        lo = max(0, int(w["x0"] / bin_w))
        hi = min(n_bins - 1, int(w["x1"] / bin_w))
        for b in range(lo, hi + 1):
            coverage[b] += 1

    total = len(words)
    lo_bin = int(page_width * 0.25 / bin_w)
    hi_bin = int(page_width * 0.75 / bin_w)
    # Find widest low-coverage band in the central region
    best_start = best_len = cur_start = cur_len = 0
    for b in range(lo_bin, hi_bin + 1):
        if coverage[b] <= max(2, total * 0.02):
            if cur_len == 0:
                cur_start = b
            cur_len += 1
            if cur_len > best_len:
                best_start, best_len = cur_start, cur_len
        else:
            cur_len = 0

    if best_len * bin_w < 14.0:  # no meaningful gutter
        return [words]

    gutter_x = (best_start + best_len / 2.0) * bin_w
    left = [w for w in words if w["x1"] <= gutter_x]
    right = [w for w in words if w["x0"] >= gutter_x]
    crossing = [w for w in words if w["x0"] < gutter_x < w["x1"]]

    if len(left) < total * 0.2 or len(right) < total * 0.2:
        return [words]

    # Both sides must span a substantial vertical extent (true sidebar)
    def v_span(ws: list[dict]) -> float:
        tops = [w["top"] for w in ws]
        return (max(tops) - min(tops)) if ws else 0.0

    span = max(v_span(words), 1.0)
    if v_span(left) < span * 0.4 or v_span(right) < span * 0.4:
        return [words]

    # Full-width lines (e.g. the name header) go with the left column
    return [crossing + left, right]


def _flow_lines(words: list[dict]) -> list[_Line]:
    if not words:
        return []
    ordered = sorted(words, key=lambda w: (w["top"], w["x0"]))

    rows: list[list[dict]] = []
    for w in ordered:
        size = w.get("size") or 10.0
        tol = max(2.0, size * 0.4)
        if rows and abs(w["top"] - rows[-1][0]["top"]) <= tol:
            rows[-1].append(w)
        else:
            rows.append([w])

    lines: list[_Line] = []
    for row in rows:
        row.sort(key=lambda w: w["x0"])
        text = _join_row(row)
        if not text.strip():
            continue
        sizes = [w["size"] for w in row if w.get("size")]
        avg = sum(sizes) / len(sizes) if sizes else None
        lines.append(_Line(text=text, top=row[0]["top"], size=avg))
    return lines


def _join_row(row: list[dict]) -> str:
    gaps = [
        row[i + 1]["x0"] - row[i]["x1"]
        for i in range(len(row) - 1)
        if row[i + 1]["x0"] > row[i]["x1"]
    ]
    positive = [g for g in gaps if 0 < g < _GAP_MIN_PT]
    typical = statistics.median(positive) if positive else 3.0
    threshold = max(_GAP_MIN_PT, typical * 4.0)

    parts: list[str] = [row[0]["text"]]
    for prev, cur in zip(row, row[1:]):
        gap = cur["x0"] - prev["x1"]
        parts.append(" | " if gap >= threshold else " ")
        parts.append(cur["text"])
    return "".join(parts)
